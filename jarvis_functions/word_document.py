from docx import Document
import win32com.client as win32
import time

from elevenlabs import play
from elevenlabs.client import ElevenLabs
import speech_recognition as sr

import os

from api_keys.api_keys import ELEVEN_LABS_API

client = ElevenLabs(api_key=ELEVEN_LABS_API)
r = sr.Recognizer()


def record_text():
    """Listen for speech and return the recognized text."""
    try:
        with sr.Microphone() as source:
            #print("Listening...")
            r.adjust_for_ambient_noise(source, duration=0.2)
            audio = r.listen(source)

            # Recognize speech using Google API
            MyText = r.recognize_google(audio, language="bg-BG")
            print(f"You said: {MyText}")
            return MyText.lower()

    except sr.RequestError as e:
        print(f"API Request Error: {e}")
        return None
    except sr.UnknownValueError:
        print("Sorry, I didn't catch that. Please try again.")
        return None

def openWord(jarvis_voice:str):
    # Initialize document
    doc = Document()

    # Open Word
    audio = client.generate(text="Разбира се, отварям Word. Само секунда", voice=jarvis_voice)
    play(audio)

    # Wait a bit to ensure Word has opened
    time.sleep(2)

    # Ask for document title
    audio = client.generate(
        text="Готов съм. Преди да започнем, как ще желаете да е заглавието на документа?",
        voice=jarvis_voice)

    play(audio)

    # Listen for the title input
    print("Listening for title...")
    input_text = record_text()

    # Add a title
    doc.add_heading(input_text, 0)

    # Inform the user
    audio = client.generate(
        text="Добре започвам да слушам и записвам. Кажете думата Край за да спра да записвам",
        voice=jarvis_voice)
    play(audio)

    words_in_document = ""

    while True:
        with sr.Microphone() as source:
            try:
                print("Listening for input...")
                input_text = record_text()
                print(f"You said: {input_text}")

                # Skip if input_text is None (empty or unrecognized speech)
                if input_text is None or input_text.strip() == "":
                    print("No speech detected or input is empty, try again.")
                    continue  # Skip to the next loop iteration

                # Stop listening when "край" is said
                if "край" in input_text or "Край" in input_text:
                    audio = client.generate(text="Спрях да записвам, файла е запазен в папка Downloads",
                                            voice=jarvis_voice)
                    play(audio)
                    # Add a paragraph
                    doc.add_paragraph(words_in_document)
                    break

                # Append the valid input to the document
                words_in_document += input_text + ". "

                time.sleep(1)  # Slight delay for realism

            except sr.UnknownValueError:
                print("Could not understand, try again.")
            except sr.RequestError:
                print("Speech recognition service error.")

    # Finished
    print("Document saved and process ended.")

    file_path = r'D:\example.docx'
    doc.save(file_path)
    os.system(f'start {file_path}')  # Open Word on Windows